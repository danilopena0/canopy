#!/usr/bin/env python3
"""Convert a markdown resume file to a formatted ATS-friendly Word document.

Usage:
    python scripts/md_to_docx.py <input.md> [output.docx]
    # Output defaults to same path with .docx extension if not specified
"""

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

BLUE = RGBColor(0x1F, 0x49, 0x7D)
LINK_COLOR = "0563C1"
_HYPERLINK_REL = "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink"

# ── Layout constants ─────────────────────────────────────────────────────────
MARGIN_TOP_BOTTOM = Inches(0.5)
MARGIN_SIDES = Inches(0.65)

FONT_NAME = 16.0       # name heading
FONT_TAGLINE = 10.0    # bold tagline under name
FONT_CONTACT = 9.0     # contact / header line
FONT_H2 = 10.0         # section headings
FONT_COMPANY = 10.0    # H3 company names
FONT_BODY = 9.5        # body text and bullets
FONT_JOBTITLE = 9.5    # bold job-title lines inside a role


def _bottom_border(paragraph, color_hex: str = "1F497D", sz: int = 4):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(sz))
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color_hex)
    pBdr.append(bottom)
    pPr.append(pBdr)


def _spacing(paragraph, before: float = 0, after: float = 0):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing_rule = WD_LINE_SPACING.SINGLE


def _add_hyperlink(paragraph, label: str, url: str, font_size: float):
    """Insert a clickable hyperlink run into a paragraph."""
    r_id = paragraph.part.relate_to(url, _HYPERLINK_REL, is_external=True)
    hl = OxmlElement("w:hyperlink")
    hl.set(qn("r:id"), r_id)
    r = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), LINK_COLOR)
    rPr.append(color)
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(u)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), str(int(font_size * 2)))
    rPr.append(sz)
    r.append(rPr)
    t = OxmlElement("w:t")
    t.text = label
    r.append(t)
    hl.append(r)
    paragraph._p.append(hl)


def _parse_inline(paragraph, text: str, base_size: float = FONT_BODY):
    """Add runs handling **bold** markers and [label](url) hyperlinks."""
    token_re = re.compile(r"\*\*(.+?)\*\*|\[([^\]]+)\]\(([^\)]+)\)")
    last = 0
    for m in token_re.finditer(text):
        if m.start() > last:
            run = paragraph.add_run(text[last:m.start()])
            run.font.size = Pt(base_size)
        if m.group(1) is not None:
            run = paragraph.add_run(m.group(1))
            run.font.size = Pt(base_size)
            run.bold = True
        else:
            _add_hyperlink(paragraph, m.group(2), m.group(3), base_size)
        last = m.end()
    if last < len(text):
        run = paragraph.add_run(text[last:])
        run.font.size = Pt(base_size)


def convert(input_path: Path, output_path: Path) -> None:
    doc = Document()

    for sec in doc.sections:
        sec.top_margin = MARGIN_TOP_BOTTOM
        sec.bottom_margin = MARGIN_TOP_BOTTOM
        sec.left_margin = MARGIN_SIDES
        sec.right_margin = MARGIN_SIDES

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(FONT_BODY)

    raw_lines = input_path.read_text(encoding="utf-8").splitlines()

    # Strip the optional agent metadata block (# Resume — ... up to first ---)
    start = 0
    if raw_lines and raw_lines[0].startswith("# Resume"):
        for i, ln in enumerate(raw_lines):
            if ln.strip() == "---":
                start = i + 1
                break

    lines = raw_lines[start:]

    after_name = False
    in_header = True

    i = 0
    while i < len(lines):
        raw = lines[i]
        line = raw.strip()
        i += 1

        if not line:
            continue

        # ── Horizontal rule ─────────────────────────────────────────────────
        if line == "---":
            p = doc.add_paragraph()
            _spacing(p, before=0, after=2)
            _bottom_border(p, color_hex="CCCCCC", sz=4)
            after_name = False
            continue

        # ── H1 = Name ───────────────────────────────────────────────────────
        if re.match(r"^# (?!#)", line):
            name = line[2:].strip()
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _spacing(p, before=0, after=1)
            run = p.add_run(name)
            run.bold = True
            run.font.size = Pt(FONT_NAME)
            after_name = True
            continue

        # ── H2 = Section heading ────────────────────────────────────────────
        if line.startswith("## "):
            heading = line[3:].strip().upper()
            p = doc.add_paragraph()
            _spacing(p, before=5, after=1)
            run = p.add_run(heading)
            run.bold = True
            run.font.size = Pt(FONT_H2)
            run.font.color.rgb = BLUE
            _bottom_border(p, color_hex="1F497D", sz=4)
            after_name = False
            in_header = False
            continue

        # ── H3 = Company / institution ──────────────────────────────────────
        if line.startswith("### "):
            company = line[4:].strip()
            p = doc.add_paragraph()
            _spacing(p, before=4, after=0)
            run = p.add_run(company)
            run.bold = True
            run.font.size = Pt(FONT_COMPANY)
            after_name = False
            continue

        # ── Bullet point ────────────────────────────────────────────────────
        if line.startswith("- "):
            content = line[2:]
            p = doc.add_paragraph(style="Normal")
            _spacing(p, before=0, after=0)
            # Hanging indent: bullet character sits at the left margin (0"),
            # text starts at 0.13" — same as a hanging indent.
            # Using Normal (not List Bullet) so Word's numPr hanging offset
            # doesn't push the bullet past the margin.
            p.paragraph_format.left_indent = Inches(0.13)
            p.paragraph_format.first_line_indent = Inches(-0.13)
            run = p.add_run("•  ")
            run.font.size = Pt(FONT_BODY)
            _parse_inline(p, content, base_size=FONT_BODY)
            after_name = False
            continue

        # ── Bold-only line (tagline or standalone bold) ─────────────────────
        m = re.match(r"^\*\*(.+)\*\*$", line)
        if m and line.count("**") == 2:
            content = m.group(1)
            p = doc.add_paragraph()
            if after_name:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                _spacing(p, before=0, after=1)
                run = p.add_run(content)
                run.bold = True
                run.font.size = Pt(FONT_TAGLINE)
                after_name = False
            else:
                _spacing(p, before=0, after=0)
                run = p.add_run(content)
                run.bold = True
                run.font.size = Pt(FONT_JOBTITLE)
            continue

        # ── Line starting with ** (job title / competency line) ─────────────
        if line.startswith("**"):
            p = doc.add_paragraph()
            _spacing(p, before=0, after=0)
            _parse_inline(p, line, base_size=FONT_BODY)
            after_name = False
            continue

        # ── Plain text ───────────────────────────────────────────────────────
        p = doc.add_paragraph()
        if in_header:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _spacing(p, before=0, after=2)
            _parse_inline(p, line, base_size=FONT_CONTACT)
        else:
            _spacing(p, before=1, after=1)
            _parse_inline(p, line, base_size=FONT_BODY)
        after_name = False

    doc.save(output_path)
    print(f"Saved: {output_path}")


FONT_CL_NAME = 14.0
FONT_CL_CONTACT = 9.5
FONT_CL_BODY = 11.0


def convert_cover_letter(input_path: Path, output_path: Path) -> None:
    """Convert a cover letter markdown file to a nicely formatted Word document.

    Strips the ATS coverage table (everything from '## ATS Coverage' onward).
    Structure expected:
        - Header lines (name + contact) before the first '---'
        - Body paragraphs after the first '---'
        - Second '---' or '## ATS Coverage' signals end of letter content
    """
    doc = Document()

    for sec in doc.sections:
        sec.top_margin = Inches(1.0)
        sec.bottom_margin = Inches(1.0)
        sec.left_margin = Inches(1.25)
        sec.right_margin = Inches(1.25)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(FONT_CL_BODY)

    raw_lines = input_path.read_text(encoding="utf-8").splitlines()

    # Strip ATS coverage block
    body_lines = []
    for ln in raw_lines:
        if ln.strip().startswith("## ATS Coverage"):
            break
        body_lines.append(ln)

    # Remove trailing '---' left by the ATS divider
    while body_lines and body_lines[-1].strip() in ("---", ""):
        body_lines.pop()

    in_header = True
    header_first_line = True
    closing_seen = False

    for line in body_lines:
        stripped = line.strip()

        # First '---' ends the header block
        if stripped == "---" and in_header:
            in_header = False
            p = doc.add_paragraph()
            _spacing(p, before=0, after=6)
            continue

        if in_header:
            if not stripped:
                continue
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if header_first_line:
                _spacing(p, before=0, after=2)
                run = p.add_run(stripped)
                run.bold = True
                run.font.size = Pt(FONT_CL_NAME)
                header_first_line = False
            else:
                _spacing(p, before=0, after=1)
                _parse_inline(p, stripped, base_size=FONT_CL_CONTACT)
            continue

        # Body
        if not stripped:
            continue

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT

        is_closing = stripped.lower().startswith("sincerely") or stripped.lower().startswith("best regards")
        is_signature = closing_seen and not stripped.lower().startswith("sincerely")

        if is_closing:
            _spacing(p, before=12, after=0)
            closing_seen = True
            run = p.add_run(stripped)
            run.font.size = Pt(FONT_CL_BODY)
        elif is_signature:
            _spacing(p, before=0, after=0)
            run = p.add_run(stripped)
            run.bold = True
            run.font.size = Pt(FONT_CL_BODY)
        else:
            _spacing(p, before=0, after=6)
            _parse_inline(p, stripped, base_size=FONT_CL_BODY)

    doc.save(output_path)
    print(f"Saved: {output_path}")


if __name__ == "__main__":
    import argparse

    parser = argparse.ArgumentParser(description="Convert markdown to Word (.docx)")
    parser.add_argument("input", help="Input .md file")
    parser.add_argument("output", nargs="?", help="Output .docx file (defaults to same path)")
    parser.add_argument("--cover", action="store_true", help="Format as a cover letter (not a resume)")
    args = parser.parse_args()

    inp = Path(args.input)
    out = Path(args.output) if args.output else inp.with_suffix(".docx")

    if args.cover:
        convert_cover_letter(inp, out)
    else:
        convert(inp, out)
